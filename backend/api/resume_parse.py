"""简历文件解析：PDF / DOCX / 纯文本 → 文本。

只用仓库里已装的依赖（pypdf、PyMuPDF(fitz)、python-docx），不新增三方包。
PDF 优先 pypdf；抽不出文字（扫描件/图片型 PDF）时回退 PyMuPDF 再试一次。
"""
from __future__ import annotations

import io

MAX_BYTES = 20 * 1024 * 1024  # 与原型一致：单文件不超过 20MB
_ALLOWED_SUFFIX = (".pdf", ".docx", ".txt", ".md")


class ResumeParseError(ValueError):
    """文件类型不支持 / 超限 / 解析不出文字。"""


def _from_pdf(data: bytes) -> str:
    text = ""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        text = ""
    if text.strip():
        return text
    # pypdf 抽不出（部分 PDF 结构或扫描件）→ 换 PyMuPDF
    try:
        import fitz

        with fitz.open(stream=data, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception:
        return ""


def _from_docx(data: bytes) -> str:
    try:
        import docx

        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        # 表格里常放技能清单，一并取出
        for t in d.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(x for x in parts if x and x.strip())
    except Exception:
        return ""


def parse_resume_bytes(filename: str, data: bytes) -> str:
    """按扩展名解析为纯文本；失败抛 ResumeParseError（含可读原因）。"""
    name = (filename or "").lower().strip()
    if not name.endswith(_ALLOWED_SUFFIX):
        raise ResumeParseError(
            f"不支持的文件类型：{filename or '(无文件名)'}；仅支持 PDF / DOCX / TXT"
        )
    if not data:
        raise ResumeParseError("文件内容为空")
    if len(data) > MAX_BYTES:
        raise ResumeParseError(
            f"文件过大：{len(data) // 1024 // 1024}MB，上限 {MAX_BYTES // 1024 // 1024}MB"
        )

    if name.endswith(".pdf"):
        text = _from_pdf(data)
    elif name.endswith(".docx"):
        text = _from_docx(data)
    else:
        text = data.decode("utf-8", errors="ignore")

    text = (text or "").strip()
    if not text:
        raise ResumeParseError(
            "未能从文件中提取到文字。若为扫描件/图片型 PDF，请改用文本粘贴方式"
        )
    return text


# 范例简历：用库内真实存在的技能名，便于一键体验时能命中技能库
SAMPLE_RESUME = """姓名：张××    求职意向：混凝土工 / 施工现场技术员

【工作经历】
2023.03 - 至今    ××建筑工程有限公司    施工班组技术员
- 负责施工前的配料准备与支模准备，核对配合比与模板尺寸；
- 熟练进行搅拌操作与泵送操作，日均浇筑量 200m³，控制坍落度在标准范围；
- 承担模板安装与拆除，配合完成场地和支泵准备；
- 每周执行设备维保与模板清理与维保，记录设备运行台账；
- 负责班组安全防护交底，参与安全隐患排查。

2021.06 - 2023.02    ××建材科技有限公司    质检员
- 负责样品检测与质量检查，出具检测报告；
- 参与生产过程管理与控制，跟踪不合格品的分析与处理。

【技能】
配料准备、支模准备、搅拌操作、泵送操作、模板安装与拆除、设备维保、
质量检查、安全隐患排查、培训与指导

【证书】
建筑施工特种作业操作证；混凝土工（四级/中级工）
"""
